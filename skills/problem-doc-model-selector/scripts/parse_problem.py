from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any


TASK_RULES = [
    ("forecasting_regression", "预测与回归", ["预测", "趋势", "估计", "回归", "未来", "时间序列", "参数估计"], ["RMSE", "MAE", "MAPE", "R2"]),
    ("classification", "分类与识别", ["分类", "判别", "识别", "风险等级", "是否", "类别"], ["Accuracy", "F1", "AUC", "混淆矩阵"]),
    ("evaluation_ranking", "综合评价与排序", ["评价", "排序", "打分", "权重", "综合指数", "择优", "排名"], ["排名稳定性", "权重敏感性", "Spearman相关"]),
    ("optimization_scheduling", "优化、调度、选址与路径", ["优化", "最小", "最大", "调度", "路径", "选址", "资源分配", "约束", "规划", "成本", "收益"], ["目标函数值", "约束违背率", "求解时间"]),
    ("clustering_profile", "聚类、分群与画像", ["聚类", "分群", "画像", "相似", "模式", "无标签"], ["轮廓系数", "稳定性", "群体特征解释"]),
    ("mechanism_simulation", "机理、仿真与系统动力学", ["机理", "仿真", "动力学", "微分方程", "差分方程", "情景", "演化"], ["历史拟合", "参数敏感性", "情景一致性"]),
    ("network_graph", "网络、图论与传播", ["网络", "节点", "边", "最短路", "最大流", "社区", "传播", "韧性"], ["连通性", "网络效率", "传播范围"]),
]

EXTERNAL_DATA_HINTS = ["收集", "搜集", "查询", "公开数据", "外部数据", "网络数据", "统计年鉴", "气象", "人口", "GDP", "经济", "交通", "遥感"]
DATA_EXTS = {".csv", ".xlsx", ".xls", ".tsv", ".txt"}
TEXT_EXTS = {".txt", ".md", ".csv", ".tsv"}
DOC_EXTS = {".docx"}
PDF_EXTS = {".pdf"}


def read_txt(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding, errors="ignore")
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return f"[DOCX 未解析，缺少 python-docx: {path.name}]"
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        return "\n".join(page.get_text("text") for page in doc)
    except Exception:
        pass
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception:
        return f"[PDF 未解析，缺少 PyMuPDF/pdfplumber: {path.name}]"


def collect_files(root: Path) -> list[Path]:
    problem_dir = root / "problem_files"
    if not problem_dir.exists():
        return []
    files = []
    for path in problem_dir.rglob("*"):
        if path.is_file() and not path.name.startswith("~"):
            files.append(path)
    return sorted(files)


def extract_text(files: list[Path]) -> tuple[str, list[dict[str, Any]]]:
    parts = []
    metadata = []
    for path in files:
        suffix = path.suffix.lower()
        if suffix in TEXT_EXTS:
            text = read_txt(path)
        elif suffix in DOC_EXTS:
            text = read_docx(path)
        elif suffix in PDF_EXTS:
            text = read_pdf(path)
        else:
            text = ""
        metadata.append({"path": str(path), "name": path.name, "suffix": suffix, "chars": len(text)})
        if text.strip():
            parts.append(f"\n\n===== {path.name} =====\n{text}")
    return "\n".join(parts), metadata


def split_questions(text: str) -> list[dict[str, Any]]:
    pattern = re.compile(r"(问题\s*[一二三四五六七八九十\d]+|第\s*[一二三四五六七八九十\d]+\s*问|Question\s*\d+|Q\s*\d+)", re.IGNORECASE)
    matches = list(pattern.finditer(text))
    questions = []
    if not matches:
        clean = compact(text)
        return [{"id": "Q1", "title": "整体任务", "text": clean[:1200], "source": "fallback"}] if clean else []

    for idx, match in enumerate(matches):
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        body = compact(text[start:end])
        questions.append({"id": f"Q{idx + 1}", "title": match.group(1), "text": body[:1600], "source": "detected"})
    return questions


def compact(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def infer_task(question_text: str) -> dict[str, Any]:
    ranked = []
    for task_id, label, triggers, metrics in TASK_RULES:
        hits = [word for word in triggers if word.lower() in question_text.lower()]
        if hits:
            ranked.append({"task_type": task_id, "label": label, "hits": hits, "metrics": metrics, "score": len(hits)})
    ranked.sort(key=lambda item: (-item["score"], item["label"]))
    if ranked:
        return ranked[0]
    return {"task_type": "general_modeling", "label": "通用数学建模", "hits": [], "metrics": ["误差指标", "稳定性", "可解释性"], "score": 0}


def detect_constraints(text: str) -> list[str]:
    candidates = []
    for pattern in [
        r"不超过[^。；;\n]{1,40}",
        r"不少于[^。；;\n]{1,40}",
        r"至少[^。；;\n]{1,40}",
        r"至多[^。；;\n]{1,40}",
        r"约束[^。；;\n]{1,60}",
        r"满足[^。；;\n]{1,60}",
        r"限制[^。；;\n]{1,60}",
        r"成本[^。；;\n]{1,60}",
        r"预算[^。；;\n]{1,60}",
        r"容量[^。；;\n]{1,60}",
    ]:
        candidates.extend(re.findall(pattern, text))
    seen = []
    for item in candidates:
        clean = item.strip()
        if clean and clean not in seen:
            seen.append(clean)
    return seen[:10]


def data_files_summary(files: list[Path]) -> list[dict[str, Any]]:
    rows = []
    for path in files:
        if path.suffix.lower() in DATA_EXTS:
            rows.append({"name": path.name, "suffix": path.suffix.lower(), "role": "attachment_or_problem_text"})
    return rows


def external_requirements(text: str) -> dict[str, Any]:
    hits = [hint for hint in EXTERNAL_DATA_HINTS if hint.lower() in text.lower()]
    tasks = []
    if hits:
        tasks.append(
            {
                "type": "manual_search",
                "query": "；".join(sorted(set(hits))) + " 相关权威公开数据",
                "notes": "由赛题解析器根据外部数据关键词自动生成，请补充具体来源 URL 或统计口径。",
                "active": True,
            }
        )
    return {"tasks": tasks}


def build_outputs(root: Path, text: str, files_meta: list[dict[str, Any]], questions: list[dict[str, Any]], files: list[Path]) -> dict[str, Any]:
    question_map = []
    for question in questions:
        task = infer_task(question["text"])
        question_map.append(
            {
                "id": question["id"],
                "title": question["title"],
                "task_type": task["task_type"],
                "task_label": task["label"],
                "matched_keywords": task["hits"],
                "suggested_metrics": task["metrics"],
                "constraints": detect_constraints(question["text"]),
                "text_excerpt": question["text"][:800],
                "expected_output": "按题面要求给出可量化结果、图表证据与逐问结论。",
                "validation": "基线对照、指标检验、敏感性/鲁棒性分析。",
            }
        )

    return {
        "project_root": str(root),
        "problem_files": files_meta,
        "data_files": data_files_summary(files),
        "questions": question_map,
        "external_data": external_requirements(text),
        "text_chars": len(text),
    }


def write_markdown(output: dict[str, Any], path: Path) -> None:
    lines = ["# 赛题解析报告", ""]
    lines.append(f"- 项目目录：`{output['project_root']}`")
    lines.append(f"- 原始文本长度：{output['text_chars']} 字符")
    lines.append(f"- 识别文件数：{len(output['problem_files'])}")
    lines.append("")
    lines.append("## 文件清单")
    for item in output["problem_files"]:
        lines.append(f"- `{item['name']}` ({item['suffix']}, {item['chars']} chars)")
    lines.append("")
    lines.append("## 子问题映射")
    for q in output["questions"]:
        lines.extend(
            [
                f"### {q['id']} {q['title']}",
                f"- 任务类型：{q['task_label']}",
                f"- 匹配关键词：{', '.join(q['matched_keywords']) or '无，按通用建模处理'}",
                f"- 建议指标：{', '.join(q['suggested_metrics'])}",
                f"- 关键约束：{'; '.join(q['constraints']) or '未自动识别，需人工核对题面'}",
                f"- 验证方式：{q['validation']}",
                "",
                q["text_excerpt"],
                "",
            ]
        )
    lines.append("## 外部数据需求")
    tasks = output["external_data"].get("tasks", [])
    if tasks:
        for task in tasks:
            lines.append(f"- {task['type']}: {task.get('query', '')} ({task.get('notes', '')})")
    else:
        lines.append("- 未自动识别强外部数据需求。")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Parse math-modeling problem files into structured planning artifacts.")
    parser.add_argument("--project", type=Path, default=Path.cwd(), help="项目根目录，默认当前目录。")
    parser.add_argument("--write-data-requirements", action="store_true", help="若识别到外部数据需求，则写入根目录 data_requirements.json。")
    args = parser.parse_args()

    root = args.project.resolve()
    files = collect_files(root)
    output_dir = root / "paper_output" / "step1"
    output_dir.mkdir(parents=True, exist_ok=True)

    text, files_meta = extract_text(files)
    questions = split_questions(text)
    output = build_outputs(root, text, files_meta, questions, files)

    (output_dir / "question_map.json").write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown(output, output_dir / "problem_analysis.md")

    data_requirements = output["external_data"]
    if args.write_data_requirements and data_requirements.get("tasks"):
        (root / "data_requirements.json").write_text(json.dumps(data_requirements, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"✅ 赛题解析已生成：{output_dir / 'problem_analysis.md'}")
    print(f"✅ 问题映射已生成：{output_dir / 'question_map.json'}")
    if args.write_data_requirements and data_requirements.get("tasks"):
        print(f"✅ 外部数据需求已生成：{root / 'data_requirements.json'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
